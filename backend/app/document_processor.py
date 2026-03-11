"""Document processing: load, chunk, embed, and store in Postgres (pgvector)."""

import hashlib
import json
import logging
from pathlib import Path
from uuid import UUID

from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import text

from app.database import get_connection

logger = logging.getLogger(__name__)

# Supported extensions and approximate max size (bytes) per file
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls", ".md"}
MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MB

EMBEDDING_DIM = 1536  # text-embedding-3-small


def _file_hash(path: Path) -> str:
    """Return SHA256 hash of file contents for change detection."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _load_text(path: Path) -> str:
    """Load raw text from a file based on extension."""
    path = path.resolve()
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md", ".csv"):
        raw = path.read_bytes()
        for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            parts = []
            for i, p in enumerate(reader.pages):
                try:
                    t = p.extract_text()
                except AttributeError as e:
                    if "get_extracted_text" in str(e):
                        logger.debug("pypdf page %s: %s (use pypdf>=4 for extract_text)", path.name, e)
                    t = ""
                except Exception as e:
                    logger.warning("pypdf page %s page %s: %s", path.name, i + 1, e)
                    t = ""
                parts.append((t or "").strip() if t else "")
            return "\n\n".join(parts).strip()
        except Exception as e:
            logger.warning("pypdf failed for %s: %s", path, e)
            return ""

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            return "\n\n".join(parts) if parts else ""
        except Exception as e:
            logger.warning("python-docx failed for %s: %s", path, e)
            return ""

    if suffix == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)
        except Exception as e:
            logger.warning("openpyxl failed for %s: %s", path, e)
            return ""

    if suffix == ".xls":
        try:
            import xlrd
            with xlrd.open_workbook(path) as wb:
                parts = []
                for sheet in wb.sheets():
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        parts.append("\t".join(str(c) if c else "" for c in row))
                return "\n".join(parts)
        except Exception as e:
            logger.warning("xlrd failed for %s: %s", path, e)
            return ""

    return ""


def _tiktoken_len(text: str) -> int:
    """Approximate token count for chunk sizing."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except Exception:
        return len(text) // 4  # fallback


def process_file(file_path: Path, embeddings: OpenAIEmbeddings) -> tuple[UUID | None, int, str]:
    """
    Load, chunk, embed, and store one file. Returns (file_id, chunk_count, status).
    Status is "ready", "error", or "processing".
    """
    path = Path(file_path).resolve()
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        return None, 0, "error"
    if path.stat().st_size > MAX_FILE_BYTES:
        return None, 0, "error"

    file_hash_val = _file_hash(path)
    filename = path.name
    file_type = path.suffix.lower().lstrip(".")
    file_size = path.stat().st_size

    with get_connection() as conn:
        # Check if already processed with same hash (optional: store hash in files)
        existing = conn.execute(
            text(
                "SELECT id, chunk_count FROM files WHERE filename = :fn AND file_size = :fs LIMIT 1"
            ),
            {"fn": filename, "fs": file_size},
        ).fetchone()
    with get_connection() as conn:
        # Re-process: delete children then parents (FKs), then re-ingest
        if existing:
            file_id = existing[0]
            conn.execute(text("DELETE FROM documents WHERE file_id = :id"), {"id": str(file_id)})
            conn.execute(text("DELETE FROM parent_chunks WHERE file_id = :id"), {"id": str(file_id)})
            try:
                conn.execute(
                    text(
                        "UPDATE files SET status = 'processing', chunk_count = 0, error_message = NULL WHERE id = :id"
                    ),
                    {"id": str(file_id)},
                )
            except Exception:
                conn.execute(
                    text("UPDATE files SET status = 'processing', chunk_count = 0 WHERE id = :id"),
                    {"id": str(file_id)},
                )

        if not existing:
            r = conn.execute(
                text(
                    "INSERT INTO files (filename, file_type, file_size, status, file_hash) "
                    "VALUES (:fn, :ft, :fs, 'processing', :fh) RETURNING id"
                ),
                {"fn": filename, "ft": file_type, "fs": file_size, "fh": file_hash_val},
            )
            row = r.fetchone()
            file_id = row[0] if row else None
            if not file_id:
                return None, 0, "error"

    def _set_error(msg: str) -> None:
        msg_trimmed = (msg or "Unknown error")[:500]
        with get_connection() as conn:
            try:
                conn.execute(
                    text(
                        "UPDATE files SET status = 'error', chunk_count = 0, error_message = :msg WHERE id = :id"
                    ),
                    {"id": str(file_id), "msg": msg_trimmed},
                )
            except Exception:
                conn.execute(
                    text("UPDATE files SET status = 'error', chunk_count = 0 WHERE id = :id"),
                    {"id": str(file_id)},
                )
                logger.warning("File %s error (run supabase_phase2_add_error_message.sql to see messages): %s", path.name, msg_trimmed)

    try:
        raw = _load_text(path)
        if not raw or not raw.strip():
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                hint = " PDF may be image-only (scanned) with no selectable text — try a text-based PDF or add a .txt version."
            elif suffix == ".docx":
                hint = " Document may be empty or use unsupported elements — try copying text into a .txt file."
            elif suffix in (".xlsx", ".xls"):
                hint = " Sheet may be empty or contain only images — ensure cells have text/numbers."
            else:
                hint = " Try saving as plain text (.txt) or check the file isn’t empty."
            _set_error(f"No text could be extracted from this file (empty or unsupported format).{hint}")
            return file_id, 0, "error"
    except Exception as e:
        logger.exception("Load failed for %s: %s", path, e)
        _set_error(f"Load failed: {e!s}")
        return file_id, 0, "error"

    # Phase 3: parent-child chunking — parents (2000 tok, 200 overlap), children (500 tok, 50 overlap)
    parent_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        length_function=_tiktoken_len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    child_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        length_function=_tiktoken_len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    parent_docs = parent_splitter.create_documents(
        [raw], metadatas=[{"source": filename, "file_id": str(file_id)}]
    )

    if not parent_docs:
        with get_connection() as conn:
            conn.execute(
                text("UPDATE files SET status = 'ready', chunk_count = 0 WHERE id = :id"),
                {"id": str(file_id)},
            )
        return file_id, 0, "ready"

    # Insert parents into parent_chunks and collect ids
    def _meta_to_json(m: dict) -> str:
        if not isinstance(m, dict):
            return "{}"
        out = {}
        for k, v in m.items():
            if isinstance(v, (str, int, float, bool)) or v is None:
                out[str(k)] = v
            elif isinstance(v, (list, dict)):
                try:
                    json.dumps(v)
                    out[str(k)] = v
                except (TypeError, ValueError):
                    out[str(k)] = str(v)
            else:
                out[str(k)] = str(v)
        return json.dumps(out)

    parent_ids: list[UUID] = []
    with get_connection() as conn:
        for i, doc in enumerate(parent_docs):
            r = conn.execute(
                text(
                    "INSERT INTO parent_chunks (content, metadata, file_id, chunk_index) "
                    "VALUES (:content, CAST(:metadata AS jsonb), :file_id, :idx) RETURNING id"
                ),
                {
                    "content": doc.page_content,
                    "metadata": _meta_to_json(doc.metadata),
                    "file_id": str(file_id),
                    "idx": i,
                },
            )
            row = r.fetchone()
            if row:
                parent_ids.append(row[0])

    # Build child docs per parent (for embedding and insert with parent_id)
    child_texts: list[str] = []
    child_meta_list: list[dict] = []
    child_parent_indices: list[int] = []  # index into parent_ids for each child
    for parent_idx, parent_doc in enumerate(parent_docs):
        children = child_splitter.create_documents(
            [parent_doc.page_content],
            metadatas=[{**parent_doc.metadata, "parent_index": parent_idx}],
        )
        for c in children:
            child_texts.append(c.page_content)
            child_meta_list.append(c.metadata)
            child_parent_indices.append(parent_idx)

    if not child_texts:
        with get_connection() as conn:
            conn.execute(
                text("UPDATE files SET status = 'ready', chunk_count = 0 WHERE id = :id"),
                {"id": str(file_id)},
            )
        return file_id, 0, "ready"

    try:
        vectors = embeddings.embed_documents(child_texts)
    except Exception as e:
        logger.exception("Embedding failed for %s: %s", path, e)
        _set_error(f"Embedding failed: {e!s}")
        return file_id, 0, "error"

    def _meta_to_json(m: dict) -> str:
        """Ensure metadata is JSON-serializable (e.g. no UUID objects)."""
        if not isinstance(m, dict):
            return "{}"
        out = {}
        for k, v in m.items():
            if isinstance(v, (str, int, float, bool)) or v is None:
                out[str(k)] = v
            elif isinstance(v, (list, dict)):
                try:
                    json.dumps(v)
                    out[str(k)] = v
                except (TypeError, ValueError):
                    out[str(k)] = str(v)
            else:
                out[str(k)] = str(v)
        return json.dumps(out)

    try:
        with get_connection() as conn:
            for i, (text_content, meta, vec, pidx) in enumerate(
                zip(child_texts, child_meta_list, vectors, child_parent_indices)
            ):
                if len(vec) != EMBEDDING_DIM:
                    continue
                parent_id = parent_ids[pidx] if pidx < len(parent_ids) else None
                conn.execute(
                    text(
                        "INSERT INTO documents (content, metadata, embedding, file_id, chunk_index, parent_id) "
                        "VALUES (:content, CAST(:metadata AS jsonb), CAST(:embedding AS vector), :file_id, :idx, :parent_id)"
                    ),
                    {
                        "content": text_content,
                        "metadata": _meta_to_json(meta),
                        "embedding": "[" + ",".join(map(str, vec)) + "]",
                        "file_id": str(file_id),
                        "idx": i,
                        "parent_id": str(parent_id) if parent_id else None,
                    },
                )
            try:
                conn.execute(
                    text(
                        "UPDATE files SET status = 'ready', chunk_count = :n, error_message = NULL WHERE id = :id"
                    ),
                    {"n": len(child_texts), "id": str(file_id)},
                )
            except Exception:
                conn.execute(
                    text("UPDATE files SET status = 'ready', chunk_count = :n WHERE id = :id"),
                    {"n": len(child_texts), "id": str(file_id)},
                )
    except Exception as e:
        logger.exception("Failed to save chunks for %s: %s", path, e)
        _set_error(f"Failed to save chunks: {e!s}")
        return file_id, 0, "error"

    return file_id, len(child_texts), "ready"


def remove_file_by_path(file_path: Path) -> None:
    """Remove file record and its chunks from DB (e.g. on file delete)."""
    path = Path(file_path).resolve()
    filename = path.name
    with get_connection() as conn:
        conn.execute(
            text("DELETE FROM documents WHERE file_id IN (SELECT id FROM files WHERE filename = :fn)"),
            {"fn": filename},
        )
        conn.execute(
            text("DELETE FROM parent_chunks WHERE file_id IN (SELECT id FROM files WHERE filename = :fn)"),
            {"fn": filename},
        )
        conn.execute(text("DELETE FROM files WHERE filename = :fn"), {"fn": filename})
